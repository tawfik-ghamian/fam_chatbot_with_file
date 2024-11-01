import re
import asyncio
import oracledb
import pdfplumber
import streamlit as st

from tqdm import tqdm
from docx import Document 
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from llama_index.core.llms import ChatMessage
from chatbot_helper import send_chatbot_request


def generate_llm(text):
    resp =  asyncio.run(send_chatbot_request(text,st.secrets["COHERE_API_KEY"])) 
    
    st.session_state.messages.append(ChatMessage(role= "assistant", content= resp))

def extract_pdf(file):
    """Extracts paragraphs, headings, and tables from a PDF file and returns a list of dictionaries.

    Args:
        file_path (str): The path to the PDF file.

    Returns:
        list: A list of dictionaries, each containing 'heading', 'paragraphs', and 'tables' keys.
    """

    data = []
    current_heading = None
    current_paragraph = []
    current_tables = []
    
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            # for tb in page.extract_tables():
            #     print("table")
            #     print(tb)
            for obj in page.filter(lambda obj: obj["object_type"] == "char").extract_text_lines():
                for char_info in obj['chars']:
                    if char_info['size'] >= 13:
                    # New heading encountered, store previous paragraph and tables
                        if current_paragraph:
                            data.append( {
                                "heading":current_heading,
                                "paragraphs": " ".join(current_paragraph),
                                "tables": current_tables
                            })
                        current_paragraph = []
                        current_tables = []
                        current_heading = obj["text"]
            
                current_paragraph.append(obj["text"])
            
            # if 
            current_tables = page.extract_tables()

            # Check for last page handling (append final paragraph and tables)
            if current_paragraph and (current_heading or current_tables):
                data.append ({
                    "heading":current_heading,
                    "paragraphs": " ".join(current_paragraph),
                    "tables": current_tables
                })
                current_paragraph = []
                current_tables = []
                
    # print(data["Invoicing of Deals (Sale and Lease)"])
    
    # print(data[6])
    
    result = {
        "file_name":file.name,
        "data":data
    }
    
    st.sidebar.success("document splitted")
    return result
     
    # with pdfplumber.open(file) as pdf:
    #     # page.filter(lambda obj: (obj["object_type"] == "char" and obj["size"] >= 12))
    #     for page in pdf.pages:
    #         for line in page:
    #             is_heading = False
    #             if line["object_type"] == "char" and line["size"] >= 12:
    #                 current_heading = line.extract_text()
    #                 print(current_heading)
    #                 is_heading = True
    #                 break
                
    #             if is_heading:
    #                 if current_paragraph:
    #                     print("heading")
    #                     print(current_heading)
    #                     data.append({
    #                         "heading": current_heading,
    #                         "paragraphs": " ".join(current_paragraph),
    #                         "tables": current_tables
    #                     })
    #                 current_heading = line
    #                 current_paragraph = []
    #                 current_tables = []
    #             else:
    #                 print("paragraph")
    #                 current_paragraph.append(line.extract_text())

    #         # Add the last paragraph and tables
    #         data.append({
    #             "heading": current_heading,
    #             "paragraphs": " ".join(current_paragraph),
    #             "tables": current_tables
    #         })

    #     return data
        
        # text = textract.process()
        # print(re.split('\s{4,}',text))
        
        # for page in pdf.pages:
        #     heading_text = page.filter(lambda obj: (obj["object_type"] == "char" and obj["size"] >= 12))
        #     # print(heading_text.extract_text())
        #     for line in heading_text.extract_text():
        #         print(line)
            #     if (line.endswith(('.', ':')) and re.match(r'^[A-Z][a-z]+\s[A-Z][a-z]+', line)) or \
            #        (line.startswith(' ') and line.endswith('.')):
            #         # New heading found
            #         if current_paragraph:
            #             data.append({
            #                 "heading": current_heading,
            #                 "paragraphs": " ".join(current_paragraph),
            #                 "tables": current_tables
            #             })
            #         current_heading = line
            #         current_paragraph = []
            #         current_tables = []
            #     else:
            #         current_paragraph.append(line)
            # if page.extract_tables():
            #     current_tables = page.extract_tables()
            #     print(current_tables)

        # Add the last paragraph and tables
        # data.append({
        #     "heading": current_heading,
        #     "paragraphs": " ".join(current_paragraph),
        #     "tables": current_tables
        # })
    
    # st.success("pdf splitted")
    # result = {
    #     "file_name":file.name,
    #     "data":data
    # }
    # return result

def extract_paragraph_docx(file):
    # doc = Document(file)
    # contents = []
    # # current_page = 1
    # current_section = {
    #     "heading": "",
    #     "paragraphs": [],
    #     "table": []
    # }
    
    # print("doc itr inner content")
    

    # for ind, k in enumerate(doc.iter_inner_content()):
    #     if ind == 0 or ind ==1:
    #         print(ind)
    #         print(k.text)
    #         print(k.style.name)
    #         print(k.style.next_paragraph_style.name)
    #         print(k.style.font.size)
        
    #     if type(k) is Table:
    #         for l in k.rows:
    #             current_section["table"].append(tuple(r.text for r in l.cells))
    #     else:
    #         # print(k.style.name) 
    #         # print(k.style.next_paragraph_style.name) 
    #         # print(k.text)
    #         if k.style.name == "Title":
    #             title = k.text
            
    #         if k.style.name.startswith("Heading") :    
                # current_section = {
                #     "heading": k.text,
                #     "paragraphs": [],
                #     "table": []
                # }
    
    #         else:
    #             if k.text != "" and k.style.name != "Title":
    #                 current_section["paragraphs"].append(k.text)
    
        # if current_section["heading"] and (len(current_section["paragraphs"]) > 1 or current_section["table"]):    
        #     contents.append(current_section)
            
    #     # print("-"*40) 
    # for ll in contents:
    #     print("00"*40)
    #     print(ll)
    # result = {
    #     "file_name": file.name,
    #     "title":title,
    #     "data": contents,
        
    # }
    # return result
    # print("doc iter_inner_content()")
    # for i in doc.iter_inner_content():
    #     for j,k in i:
    #         print(j)
    #         print(k)

    # for paragraph in tqdm(doc.paragraphs):
    #     # Check if new page
        
    #     print("iter_inner_content()")
    #     for i in paragraph.iter_inner_content():
    #         print(i)
            
    #     if paragraph.page.page_number > current_page:
            
    #         current_page += 1
    #         # Append previous section if it has content
    #         if current_section["heading"] or current_section["paragraphs"] or current_section["table"]:
    #             contents.append(current_section)
    #         # Reset for new page
    #         current_section = {
    #             "page": current_page,
    #             "heading": "",
    #             "paragraphs": [],
    #             "table": []
    #         }

    #     # Check for heading based on a combination of style and font size (optional)
    #     is_heading = paragraph.style.name.startswith("Heading") and paragraph.style.font.size > 12

    #     if is_heading:
    #         current_section["heading"] = paragraph.text

    #     # Check for normal paragraph
    #     else:
    #         current_section["paragraphs"].append(paragraph.text)

    #     # Check for tables
    #     for table in paragraph.tables:
    #         table_data = []
    #         for row in table.rows:
    #             row_data = []
    #             for cell in row.cells:
    #                 row_data.append(cell.text)
    #             table_data.append(row_data)

    #         current_section["table"] = table_data

    # # Append the last section
    # if current_section["heading"] or current_section["paragraphs"] or current_section["table"]:
    #     contents.append(current_section)

    # result = {
    #     "file_name": file.name,
    #     "data": contents
    # }
    
    # for i in contents:
    #     print("--"*50+"\n")
    #     print(i)
    #     print("--"*50+"\n")
    
    # return result
    
    doc = Document(file)
    contents = []
    paras = []
    data = {
        "heading": "",
        "paragraphs": [],
        "table": []
    }
    header = ""
    
    for element in tqdm(doc.element.body):
        if isinstance(element, CT_Tbl):
            table = Table(element, doc)
            rows = []
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)
                rows.append(tuple(text))
            data["table"] = rows  # Store table data
            continue
        
        if isinstance(element, CT_P):
            paragraph = Paragraph(element, doc)
            if paragraph.style != None and paragraph.style.name.startswith('Heading') and paragraph.style.name in ['Heading 1','Heading 2', 'Heading 3']:
                header_temp = paragraph.text.strip()
                if header_temp:  # Check if header is not empty
                    # Append data only if it has content to avoid duplicates
                    if data["heading"] and (data["paragraphs"] or data["table"]):
                        contents.append(data)
                    
                    # Reset for the new section
                    header = header_temp
                    data = {
                        "heading": header,
                        "paragraphs": [],
                        "table": []
                    }
                    paras = []
            elif paragraph.text.strip() and paragraph.text != header:  # Avoid empty and header texts
                paras.append(paragraph.text.strip())

            data["paragraphs"] = paras  # Update paragraphs in data

    # Append the last section if it has content
    if data["heading"] and (data["paragraphs"] or data["table"]):
        contents.append(data)
    
    # for ll in contents:
    #     print("00"*40)
    #     print(ll)
    formatted = []
    for j in contents:
        if len(j["paragraphs"]) > 10:
            for i in range(0, len(j["paragraphs"]), 10):
                splitted = j["paragraphs"][i:i+10]
                formatted.append({
                    "heading":j["heading"],
                    "paragraphs":splitted,
                    "table":j["table"]
                })
        else:
            formatted.append(j)
                    
    result = {
        "file_name":file.name,
        "data":formatted
    }
    
    print(result["file_name"])
    for k in result["data"]:
        print("-"*40)
        print(k)
    
    st.sidebar.success("document splitted")
    return result

def insert_paragraph_to_oracledb(contents):
    connection = oracledb.connect(user="ai", password="testtest",dsn="91.75.21.131:9522/FREEPDB1")
    cursor = connection.cursor()
    # insert into my_vectors(id,file_name,title,extracted_text,vectors)
    for ind,content in enumerate(contents["data"]):
        embedding_text = []
        embedding_text.append(content["heading"])
        embedding_text.append('\n')
        for i in content["paragraphs"]:
            embedding_text.append(i)
        embedding_text.append('\n')
        for j in content["table"]:
            embedding_text.append(j)
        embedding_text = " ".join([str(i) for i in embedding_text])
        cleaned_text = embedding_text.replace("\'" , "")
        sql_statement = cursor.execute("""
                                        insert into ai.docs(file_name,title,extracted_text,vectors)
                                        select :file_name file_name,:title title,:extracted_text extracted_text, TO_VECTOR(VECTOR_EMBEDDING(ALL_MINILM_L12_V2 USING :cleaned_text as data)) vectors
                                        from dual
                                    """,file_name=contents["file_name"],title=content["heading"],extracted_text=embedding_text,cleaned_text=cleaned_text)
        print(sql_statement)
        print("added")

    connection.commit()
    print("committed")
    cursor.close()
    connection.close()
    st.sidebar.success("The file uploaded successfully!")



st.title("🤖 Welcome in :blue[_fam_ _property_] ChatBot :sunglasses:")

@st.dialog("Sign Up")
def email_form():
    email = st.text_input("Name")
    password = st.text_input("Password")
    print(password)
    print(email)
    
    if st.button('submit') and email == "admin" and password == "admin":
        st.session_state.messages.append(ChatMessage(role= "user", content="signing in"))
        st.session_state.messages.append(ChatMessage(role= "assistant", content="Thanks for signing in. you can complete :smile:"))
        st.rerun()

if "messages" not in st.session_state:
    st.session_state.messages = [
        ChatMessage(role="system", content="You are a real state assistant that helps users find best properties in Dubai that fit there requirement")
]        

if len(st.session_state.messages) == 7:
    st.write("please sign in to fam property to complete conversation")
    st.button('Sign In',on_click=email_form)
        
else:
    uploaded_file = st.sidebar.file_uploader("Upload a file", type=["txt","docx","pdf"])
    if uploaded_file is not None:
        if uploaded_file.name.__contains__(".pdf"):
            print("pdf file")
            res = extract_pdf(uploaded_file)
            # print(res)
            insert_paragraph_to_oracledb(res)
        elif uploaded_file.name.__contains__(".docx"):
             print("docx file")
             res = extract_paragraph_docx(uploaded_file)
            #  print(res)
             insert_paragraph_to_oracledb(res)
        elif uploaded_file.name.__contains__(".txt"):
            data = []
            file_content = uploaded_file.read().decode("utf-8")
            paragraphs = file_content.split('\n\n')
            for p in paragraphs:
                data.append({
                        "heading": "",
                        "paragraphs": [p],
                        "tables": []
                    })
            
            result = {
                "file_name":uploaded_file.name,
                "data":data
            }
        
        print(uploaded_file._file_urls.upload_url)
        print(uploaded_file.name)
            
        prompt = st.chat_input("Ask a question?")
                        
        if prompt:
            
            st.session_state.messages.append(ChatMessage(role= "user", content=prompt))
            generate_llm(prompt)
        
        for message in st.session_state.messages:
            if message.role != "system":
                with st.chat_message(message.role):
                    st.markdown(message.content)